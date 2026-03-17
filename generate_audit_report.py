from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Color palette ─────────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0D, 0x1B, 0x2A)   # #0D1B2A
ACCENT_BLUE = RGBColor(0x1B, 0x6C, 0xA8)   # #1B6CA8
MID_GRAY    = RGBColor(0x55, 0x65, 0x75)   # #556575
LIGHT_GRAY  = RGBColor(0xF4, 0xF6, 0xF8)   # #F4F6F8
RED         = RGBColor(0xC0, 0x39, 0x2B)   # #C0392B
GREEN       = RGBColor(0x1A, 0x7A, 0x4A)   # #1A7A4A
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)   # #E67E22
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'),    str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color','000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_horizontal_rule(doc, color='1B6CA8'):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_score_badge(doc, score, label='Overall Score'):
    """Renders a wide table acting as a score banner."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    widths = [Inches(1.4), Inches(1.0), Inches(4.0)]
    row    = tbl.rows[0]
    for i, w in enumerate(widths):
        row.cells[i].width = w

    # Label cell
    c0 = row.cells[0]
    set_cell_bg(c0, '0D1B2A')
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(label.upper())
    r0.bold = True
    r0.font.size  = Pt(9)
    r0.font.color.rgb = WHITE
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Score cell
    c1 = row.cells[1]
    color = '1A7A4A' if score >= 60 else ('E67E22' if score >= 45 else 'C0392B')
    set_cell_bg(c1, color)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f'{score}/100')
    r1.bold = True
    r1.font.size  = Pt(14)
    r1.font.color.rgb = WHITE
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Blank spacer
    c2 = row.cells[2]
    set_cell_bg(c2, 'F4F6F8')

    doc.add_paragraph()

def style_heading(p, text, level=1):
    p.clear()
    if level == 1:
        run = p.add_run(text)
        run.bold = True
        run.font.size  = Pt(22)
        run.font.color.rgb = DARK_NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size  = Pt(10)
        run.font.color.rgb = ACCENT_BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(2)
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.font.size  = Pt(11)
        run.font.color.rgb = DARK_NAVY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

def add_h1(doc, text):
    p = doc.add_paragraph()
    style_heading(p, text, 1)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    style_heading(p, text, 2)
    return p

def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size  = Pt(10)
        rb.font.color.rgb = DARK_NAVY
    r = p.add_run(text)
    r.font.size  = Pt(10)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p

def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    # Split on bold marker **...**
    parts = text.split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.bold = (i % 2 == 1)
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p

def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        parts = item.split('**')
        for j, part in enumerate(parts):
            r = p.add_run(part)
            r.bold = (j % 2 == 1)
            r.font.size  = Pt(10)
            r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

def add_impact_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'EBF5FB')
    set_cell_border(cell,
        top={'val':'single','sz':8,'color':'1B6CA8'},
        left={'val':'single','sz':8,'color':'1B6CA8'},
        bottom={'val':'single','sz':4,'color':'1B6CA8'},
        right={'val':'single','sz':4,'color':'1B6CA8'},
    )
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(10)
    r.font.color.rgb = DARK_NAVY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '0D1B2A')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size  = Pt(9)
        r.font.color.rgb = WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri+1]
        bg   = 'FFFFFF' if ri % 2 == 0 else 'F4F6F8'
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            # Color-code Yes/No
            color = RGBColor(0x2C,0x3E,0x50)
            bold  = False
            if val in ('Yes','FAIL','Critical'):
                color = RED; bold = True
            elif val in ('No','PASS','Good'):
                color = GREEN; bold = True
            elif val.startswith('Partial') or val.startswith('©'):
                color = ORANGE; bold = True
            r = p.add_run(val)
            r.font.size  = Pt(9)
            r.font.color.rgb = color
            r.bold = bold

    if col_widths:
        for ri2 in range(len(tbl.rows)):
            for ci2, w in enumerate(col_widths):
                tbl.rows[ri2].cells[ci2].width = w

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(60)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('WEBSITE AUDIT REPORT')
r.bold = True
r.font.size  = Pt(28)
r.font.color.rgb = DARK_NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Complete Coach Works Family of Companies')
r2.font.size  = Pt(16)
r2.font.color.rgb = ACCENT_BLUE
r2.bold = True

doc.add_paragraph()
add_horizontal_rule(doc, '1B6CA8')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta.add_run(f'Prepared for: Dale Carson\nDate: {datetime.date.today().strftime("%B %d, %Y")}\nPrepared by: NuStack Digital Ventures')
r3.font.size  = Pt(11)
r3.font.color.rgb = MID_GRAY

doc.add_paragraph()
doc.add_paragraph()

sites_tbl = doc.add_table(rows=4, cols=3)
sites_tbl.style = 'Table Grid'
sites_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_c = ['Site', 'URL', 'Score']
widths_c  = [Inches(1.8), Inches(3.2), Inches(1.0)]
hrow_c    = sites_tbl.rows[0]
for i, h in enumerate(headers_c):
    cell = hrow_c.cells[i]
    cell.width = widths_c[i]
    set_cell_bg(cell, '0D1B2A')
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE

cover_data = [
    ('Complete Coach Works', 'completecoach.com', '48/100', 'C0392B'),
    ('Shuttle Bus Leasing',  'sblbus.com',         '41/100', 'C0392B'),
    ('Transit Sales Intl',   'transitsales.com',   '52/100', 'E67E22'),
]
for ri, (name, url, score, clr) in enumerate(cover_data):
    row = sites_tbl.rows[ri+1]
    bg  = 'FFFFFF' if ri % 2 == 0 else 'F4F6F8'
    for ci, val in enumerate([name, url, score]):
        cell = row.cells[ci]
        cell.width = widths_c[ci]
        set_cell_bg(cell, bg if ci < 2 else clr)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 2 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size  = Pt(10)
        r.bold = (ci == 2)
        r.font.color.rgb = WHITE if ci == 2 else RGBColor(0x2C,0x3E,0x50)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SITE 1 — completecoach.com
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'Site 1 — completecoach.com')
add_body(doc, 'Complete Coach Works  |  Platform: WordPress + Divi 4.27.6  |  Riverside, CA')
add_horizontal_rule(doc)
add_score_badge(doc, 48)

add_h2(doc, 'Critical Issues')
issues_ccw = [
    '**No crawlable content** — Divi renders everything in JavaScript. Google sees a nearly blank page. The company claiming "nation\'s largest" is invisible in search.',
    '**No sitemap** — /sitemap.xml returns 404. A "Coming Soon" zombie page is live and indexed by Google.',
    '**No meta description** on the homepage — Google writes its own snippet, typically irrelevant text.',
    '**No schema markup** — competitors get rich results (address, phone, hours) in Google search. CCW gets nothing.',
    '**Phone number and email buried** — no contact info above the fold on a site selling $500K+ service contracts.',
    '**Blog is active but disconnected** — APTA award win, RATP Dev contract, new President appointment all generating zero search traffic because the content is not properly structured.',
    '**Only 6 discoverable pages** — no services hub, no case studies, no fleet type pages, no RFQ flow.',
]
for item in issues_ccw:
    add_bullet(doc, item)

add_h2(doc, 'Revenue Impact')
add_impact_box(doc,
    'Conservative opportunity cost: $200K–$800K/year in lost inbound contracts.\n'
    'A transit agency procurement manager searching "bus refurbishment contractor" or "ZEPS electric bus conversion" will not find CCW. '
    'Three major press events in early 2026 drove zero search traffic because the site cannot be crawled by Google.'
)

add_h2(doc, 'Automation Opportunities')
auto_ccw = [
    '**RFQ pipeline** — intake form (fleet size, bus type, service needed, timeline) → auto-assign to regional sales rep → CRM entry + email confirmation',
    '**Press/contract syndication** — new contract wins auto-post to LinkedIn + trigger email blast to agency procurement contacts',
    '**Parts inquiry routing** — self-serve parts request form with auto-acknowledgment eliminates inbound phone volume',
]
for item in auto_ccw:
    add_bullet(doc, item)

add_h2(doc, 'Top 5 Fixes — Priority Order')
fixes_ccw = [
    '**Fix crawlability immediately** — generate and submit /wp-sitemap.xml, delete the "Coming Soon" zombie page',
    '**Add meta descriptions** to every page (homepage, services, about, contact minimum)',
    '**Implement Organization + LocalBusiness schema** with phone, address, and social links',
    '**Put phone number and email in the header** on every page — $0 fix with direct revenue impact',
    '**Build a proper Services architecture** — individual pages per service line with target keywords and a quote CTA',
]
add_numbered(doc, fixes_ccw)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SITE 2 — sblbus.com
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'Site 2 — sblbus.com')
add_body(doc, 'Shuttle Bus Leasing  |  Platform: WordPress + Divi 4.27.6 + WooCommerce  |  Riverside, CA')
add_horizontal_rule(doc)
add_score_badge(doc, 41)

add_h2(doc, 'Critical Issues')
issues_sbl = [
    '**No crawlable content** — same Divi JS-rendering problem as CCW. Homepage is invisible to Google.',
    '**No meta description** — confirmed missing. Google generates its own snippet, typically CSS code or nav text.',
    '**Blog last updated October 2023** — 17 months stale. Google\'s freshness algorithm treats this as a dormant business.',
    '**No pricing on inventory** — every listing forces a "Request Information" phone call. Buyers researching move on to competitors who show price ranges.',
    '**Copyright says © 2022** — it is 2026. A stale copyright year is a trust signal to first-time visitors.',
    '**No lead capture on homepage** — no form, no phone in the hero, no "Get a Lease Quote" CTA. Every visitor who bounces is gone with zero follow-up capability.',
    '**"Prison Transports" in main nav** — sitting alongside transit agency products. Brand perception issue for the primary buyer persona.',
    '**WooCommerce cart/checkout pages live and indexable** — dead infrastructure wasting Google\'s crawl budget.',
]
for item in issues_sbl:
    add_bullet(doc, item)

add_h2(doc, 'Revenue Impact')
add_impact_box(doc,
    'One transit agency on a multi-year lease is worth $50K–$500K.\n'
    'SBL ranks for essentially no lease-specific search terms — "short term bus lease," "gap fleet leasing transit," "seasonal shuttle bus lease." '
    'Zero lead capture means zero follow-up on every visitor who does not pick up the phone.'
)

add_h2(doc, 'Automation Opportunities')
auto_sbl = [
    '**Lease inquiry automation** — intake form → auto-route to sales rep → confirmation email with available inventory PDF → 3-day and 7-day follow-up sequence',
    '**Inventory availability alerts** — "notify me when a 40ft low-floor becomes available" captures buyers not ready today',
    '**Fleet gap analysis calculator** — emergency rehab coverage tool with lead capture at the end',
    '**Cross-sell routing** — auto-introduction to TSI (purchase) or CCW (rehab) when lease is not the right fit',
]
for item in auto_sbl:
    add_bullet(doc, item)

add_h2(doc, 'Top 5 Fixes — Priority Order')
fixes_sbl = [
    '**Add phone number and "Get a Quote" CTA to every page header** — 30-minute fix, immediate conversion impact',
    '**Update copyright year** (2022 → 2026) across all three sites',
    '**Add price ranges** ("starting at $X/month") to inventory listings — removes the friction killing conversion',
    '**Publish 2 blog posts/month** targeting lease-specific keywords — compounds in search within 60–90 days',
    '**Implement Organization + LocalBusiness schema** with consistent NAP across all three properties',
]
add_numbered(doc, fixes_sbl)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SITE 3 — transitsales.com
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'Site 3 — transitsales.com')
add_body(doc, 'Transit Sales International  |  Platform: WordPress + Divi 4.27.6 + WooCommerce + Yoast SEO  |  Riverside, CA')
add_horizontal_rule(doc)
add_score_badge(doc, 52)

add_h2(doc, 'Critical Issues')
issues_tsi = [
    '**Homepage title is "home - Transit Sales International"** — the word "home" leading a title tag is textbook SEO misconfiguration. Yoast is installed but not properly set up.',
    '**No H1 tag on the homepage** — confirmed absent. H1 is the single highest-weighted on-page SEO signal.',
    '**Inventory is stale** — listings dated 2019, last modified 2021. No "Available/Sold" status on any listing.',
    '**70+ listings, zero pricing** — every unit forces "Request a Quote" with no price signal. Buyers comparing dealers will not make the call.',
    '**Product images are low-quality/generic** — references to "Untitled-1.jpg" from 2019. Used vehicle buyers are visual — poor photos destroy trust.',
    '**About page last updated March 2021** — 5 years stale. New President, RATP Dev contract, TriMet BEV conversion milestone — none reflected.',
    '**Resources page slug exposes internal company structure** — routes buyers away from the conversion funnel to a page about sister companies.',
    '**WooCommerce cart live for $40K bus transactions** — no buyer is adding a bus to a shopping cart. Dead infrastructure, wasted crawl budget.',
]
for item in issues_tsi:
    add_bullet(doc, item)

add_h2(doc, 'Revenue Impact')
add_impact_box(doc,
    'TSI claims "largest inventory in the nation with over 1,000 makes and models" but only 70+ are listed.\n'
    'The other 930 units do not exist to Google or any buyer who finds the site through search.\n\n'
    'Estimated lift from fixes alone: $245K+ annually\n'
    '(70 listed units × $35K average × 10% conversion improvement from better UX and visible pricing)'
)

add_h2(doc, 'Automation Opportunities')
auto_tsi = [
    '**Inventory sync** — real-time Available/Pending/Sold status tied to internal system; auto-remove sold units',
    '**Automated quote response** — immediate reply with specs PDF, 3 comparable alternatives, and a calendar link for a sales call',
    '**Saved search alerts** — "notify me when a 35ft low-floor diesel comes in" captures buyers 60 days from decision',
    '**Fleet procurement track** — separate intake form for agencies buying 5–20 buses, routes to senior rep with different SLA',
    '**Cross-company routing** — buyer needing refurbishment before purchase gets automatic CCW introduction',
]
for item in auto_tsi:
    add_bullet(doc, item)

add_h2(doc, 'Top 5 Fixes — Priority Order')
fixes_tsi = [
    '**Fix Yoast configuration** — update homepage title to lead with brand + keyword, write a 155-character meta description, add H1 to homepage',
    '**Add pricing or price ranges** ("starting at $28,000") to all inventory listings',
    '**Audit all 70+ listings** — mark sold units, add "Date Listed," replace placeholder images with real photos, add mileage and condition',
    '**Verify sitemap in Google Search Console** — Yoast is generating it but indexation is unconfirmed (30-minute task)',
    '**Update the About page** with 2026 milestones: new President, RATP Dev contract, TriMet BEV conversion',
]
add_numbered(doc, fixes_tsi)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY PAGE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'Summary — All Three Properties')
add_horizontal_rule(doc)
doc.add_paragraph()

add_h2(doc, 'Score Overview')
add_table(doc,
    headers=['Site', 'Score', 'Biggest Single Problem'],
    rows=[
        ('completecoach.com', '48/100', 'Content invisible to Google (Divi JS rendering)'),
        ('sblbus.com',         '41/100', 'No pricing, no lead capture, blog 17 months stale'),
        ('transitsales.com',   '52/100', 'Yoast misconfigured, stale inventory, no pricing'),
    ],
    col_widths=[Inches(2.0), Inches(1.0), Inches(3.4)]
)

add_h2(doc, 'Shared Issues Across All Three Sites')
add_table(doc,
    headers=['Issue', 'CCW', 'SBL', 'TSI'],
    rows=[
        ('Missing meta description',            'Yes', 'Yes', 'Partial (17 chars)'),
        ('No H1 on homepage',                   'Yes', 'Yes', 'Yes'),
        ('Divi blocking content from Google',   'Yes', 'Yes', 'Yes'),
        ('No schema markup',                    'Yes', 'Yes', 'Partial'),
        ('Stale content / copyright',           'Yes', '© 2022', 'Last updated 2021'),
        ('No pricing visible',                  'N/A', 'Yes', 'Yes'),
        ('No lead capture on homepage',         'Yes', 'Yes', 'Yes'),
        ('Sitemap functional',                  'No',  'No',  'Yes (Yoast)'),
    ],
    col_widths=[Inches(3.0), Inches(1.0), Inches(1.0), Inches(1.4)]
)

add_h2(doc, 'Key Observations')
obs = [
    '**All three sites run Divi 4.27.6** — one developer manages all three. Every fix can be batched across all three simultaneously, reducing implementation cost.',
    '**All three share the same phone and address** — this is a local SEO asset, but only if all three have matching, correct schema markup. Currently none do.',
    '**Single highest-ROI fix available:** add a phone number and quote CTA to the header of every page on all three sites. Two hours of work. Immediate conversion impact across the entire company family.',
    '**WooCommerce is the wrong tool** for a B2B quote-and-call sales model. The cart, checkout, and my-account pages are dead infrastructure across all three sites wasting Google\'s crawl allocation.',
]
for item in obs:
    add_bullet(doc, item)

doc.add_paragraph()
add_horizontal_rule(doc, '0D1B2A')
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(f'NuStack Digital Ventures  ·  Confidential  ·  {datetime.date.today().strftime("%B %Y")}')
fr.font.size  = Pt(9)
fr.font.color.rgb = MID_GRAY

# ── Save ──────────────────────────────────────────────────────────────────────
path = r'C:\Users\bradp\dev\bus-engine\dale-carson-website-audit.docx'
doc.save(path)
print(f'Saved: {path}')
